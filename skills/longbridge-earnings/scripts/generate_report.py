#!/usr/bin/env python3
"""
Earnings Update Report — Template Script
=========================================
Usage: Fill in the DATA SECTION below with actual company data,
then run: python3 generate_report.py

Handles:
  - CJK / Latin bilingual chart labels
  - Charts embedded in DOCX via BytesIO (no PNG files saved to disk)
  - Standard 8-12 page report structure
"""

import io
import os
import sys
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─────────────────────────────────────────────
# FONT SETUP  (bilingual: CJK + Latin fallback)
# ─────────────────────────────────────────────
_CJK_FONTS = [
    'PingFang SC', 'Heiti SC', 'STHeiti', 'Songti SC',  # macOS
    'SimHei', 'Microsoft YaHei',                          # Windows
    'Noto Sans CJK SC', 'WenQuanYi Micro Hei',           # Linux
]
_available_fonts = {f.name for f in fm.fontManager.ttflist}
CHART_FONT = next((f for f in _CJK_FONTS if f in _available_fonts), 'DejaVu Sans')
matplotlib.rcParams['font.family'] = CHART_FONT
matplotlib.rcParams['axes.unicode_minus'] = False  # fix minus sign rendering


def make_chart(fig) -> io.BytesIO:
    """Save a matplotlib figure to BytesIO and return it. Do NOT save to disk."""
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close(fig)
    return buf


# ─────────────────────────────────────────────
# DATA SECTION — fill in before running
# ─────────────────────────────────────────────
SYMBOL       = "700.HK"
COMPANY      = "Tencent Holdings"
QUARTER      = "Q4 2025"
REPORT_DATE  = "2026-04-29"
RATING       = "BUY"
PRICE_TARGET = 700
CURR_PRICE   = 478.2
CURRENCY     = "HKD"

# Core KPIs: (label, reported, yoy_pct, vs_est_pct, beat=True/False)
KPIS = [
    ("Revenue (RMB B)",       194.4, 13.0,  0.5,  True),
    ("Gross Profit (RMB B)",  108.3, 19.0,  2.8,  True),
    ("Non-IFRS NI (RMB B)",    66.7, 18.0,  2.9,  True),
    ("GAAP EPS (HKD)",          6.93, 19.2, -3.8, False),
    ("Non-IFRS EPS (HKD)",      7.75, 23.2,  3.3,  True),
]

# Quarterly revenue: list of (quarter_label, revenue)
QUARTERLY_REVENUE = [
    ("Q1'24", 159.5), ("Q2'24", 161.1), ("Q3'24", 167.2), ("Q4'24", 172.4),
    ("Q1'25", 177.5), ("Q2'25", 185.9), ("Q3'25", 194.1), ("Q4'25", 194.4),
]

# Segments: list of (name, revenue, yoy_pct)
SEGMENTS = [
    ("FinTech & Biz Svcs", 60.8,  8),
    ("Marketing Svcs",     41.1, 17),
    ("Domestic Games",     38.2, 15),
    ("Social Networks",    30.6,  3),
    ("Intl Games",         21.1, 32),
]

# Quarterly margins: list of (quarter_label, gross_margin, op_margin, net_margin)
MARGINS = [
    ("Q1'24", 48.5, 30.1, 25.2),
    ("Q2'24", 49.2, 31.0, 26.0),
    ("Q3'24", 50.1, 32.5, 27.1),
    ("Q4'24", 51.3, 33.0, 27.8),
    ("Q1'25", 52.0, 33.8, 28.5),
    ("Q2'25", 53.1, 34.5, 29.0),
    ("Q3'25", 54.0, 35.2, 29.8),
    ("Q4'25", 55.7, 35.8, 30.1),
]

# Estimate revisions: list of (metric, old_val, new_val, unit)
EST_REVISIONS = [
    ("Revenue",      845, 860, "RMB B"),
    ("Non-IFRS NI",  270, 280, "RMB B"),
    ("Non-IFRS EPS", 131, 135, "HKD"),
]

# Investment thesis: list of (emoji_status, pillar, one_liner)
THESIS = [
    ("🟢", "AI Monetisation",    "Cloud +22% YoY; 3 price hikes in 12 months"),
    ("🟢", "Intl Games Breakout","USD 10B annualised; reduces China reg risk"),
    ("🟡", "Advertising",        "+17% YoY; Weixin closed-loop driving CPM"),
    ("🟠", "Margin Compression", "AI capex >2x; FY2026E op margin 36.0%"),
]

# Key risks: list of strings
RISKS = [
    "AI Capex Overrun", "Games Regulatory Risk", "WeChat Competition",
    "FinTech Deceleration", "FX Headwind", "Valuation Re-rating Delay",
]

OUTPUT_FILE = f"{SYMBOL.replace('.', '')}_{QUARTER.replace(' ', '_')}_Earnings_Update.docx"


# ─────────────────────────────────────────────
# CHART GENERATORS
# ─────────────────────────────────────────────

def chart_quarterly_revenue():
    labels = [r[0] for r in QUARTERLY_REVENUE]
    values = [r[1] for r in QUARTERLY_REVENUE]
    colors = ['#4472C4'] * (len(values) - 1) + ['#ED7D31']

    fig, ax = plt.subplots(figsize=(8, 4))
    bars = ax.bar(labels, values, color=colors)
    ax.bar_label(bars, fmt='%.1f', padding=2, fontsize=8)
    ax.set_title(f"{COMPANY} — Quarterly Revenue ({CURRENCY} B)", fontsize=11)
    ax.set_ylabel(f"Revenue ({CURRENCY} B)")
    ax.set_ylim(0, max(values) * 1.15)
    ax.tick_params(axis='x', rotation=30)
    fig.tight_layout()
    return make_chart(fig)


def chart_margins():
    labels = [r[0] for r in MARGINS]
    gross  = [r[1] for r in MARGINS]
    op     = [r[2] for r in MARGINS]
    net    = [r[3] for r in MARGINS]

    fig, ax = plt.subplots(figsize=(8, 4))
    ax.plot(labels, gross, 'o-', label='Gross Margin', color='#4472C4')
    ax.plot(labels, op,    's-', label='Op Margin',    color='#ED7D31')
    ax.plot(labels, net,   '^-', label='Net Margin',   color='#70AD47')
    ax.set_title(f"{COMPANY} — Quarterly Margin Trends (%)", fontsize=11)
    ax.set_ylabel("Margin (%)")
    ax.legend(fontsize=8)
    ax.tick_params(axis='x', rotation=30)
    fig.tight_layout()
    return make_chart(fig)


def chart_segments():
    names  = [s[0] for s in SEGMENTS]
    revs   = [s[1] for s in SEGMENTS]
    colors = ['#4472C4', '#ED7D31', '#A9D18E', '#FFC000', '#9DC3E6']

    fig, ax = plt.subplots(figsize=(6, 4))
    bars = ax.barh(names, revs, color=colors)
    ax.bar_label(bars, fmt='%.1f', padding=3, fontsize=8)
    ax.set_title(f"{COMPANY} — Revenue by Segment ({QUARTER})", fontsize=11)
    ax.set_xlabel(f"Revenue ({CURRENCY} B)")
    ax.invert_yaxis()
    fig.tight_layout()
    return make_chart(fig)


def chart_estimate_revisions():
    metrics = [r[0] for r in EST_REVISIONS]
    old_vals = [r[1] for r in EST_REVISIONS]
    new_vals = [r[2] for r in EST_REVISIONS]
    x = range(len(metrics))
    w = 0.35

    fig, ax = plt.subplots(figsize=(7, 4))
    b1 = ax.bar([i - w/2 for i in x], old_vals, w, label='Old Est', color='#A9A9A9')
    b2 = ax.bar([i + w/2 for i in x], new_vals, w, label='New Est', color='#4472C4')
    ax.bar_label(b1, fmt='%.0f', padding=2, fontsize=8)
    ax.bar_label(b2, fmt='%.0f', padding=2, fontsize=8)
    ax.set_title("FY2026E Estimate Revisions", fontsize=11)
    ax.set_xticks(list(x))
    ax.set_xticklabels(metrics)
    ax.legend()
    fig.tight_layout()
    return make_chart(fig)


# ─────────────────────────────────────────────
# DOCX BUILDER
# ─────────────────────────────────────────────

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_kpi_table(doc):
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['Metric', 'Reported', 'YoY', 'vs Estimate', 'Beat/Miss']):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True

    for kpi in KPIS:
        row = table.add_row().cells
        row[0].text = kpi[0]
        row[1].text = str(kpi[1])
        row[2].text = f"{'+' if kpi[2] >= 0 else ''}{kpi[2]:.1f}%"
        row[3].text = f"{'+' if kpi[3] >= 0 else ''}{kpi[3]:.1f}%"
        row[4].text = "Beat ✓" if kpi[4] else "Miss ✗"
    doc.add_paragraph()


def build_docx():
    doc = Document()

    # ── Page 1: Summary ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{COMPANY} ({SYMBOL})")
    run.bold = True
    run.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{QUARTER} Earnings Update  |  {REPORT_DATE}").font.size = Pt(11)

    doc.add_paragraph(
        f"Rating: {RATING}  |  Price Target: {CURRENCY} {PRICE_TARGET}  "
        f"|  Current: {CURRENCY} {CURR_PRICE}  "
        f"|  Upside: {(PRICE_TARGET/CURR_PRICE - 1)*100:+.1f}%"
    ).runs[0].bold = True

    doc.add_paragraph()
    add_heading(doc, "Key Performance Indicators", level=2)
    add_kpi_table(doc)

    # ── Pages 2-3: Revenue & Margins ──
    doc.add_page_break()
    add_heading(doc, "Quarterly Revenue Trend", level=2)
    doc.add_picture(chart_quarterly_revenue(), width=Inches(6))
    doc.add_paragraph("Source: Company filings, Longbridge CLI").italic = True  # noqa — set via run below
    # fix italic on source line
    doc.paragraphs[-1].runs[0].italic = True

    add_heading(doc, "Margin Trends", level=2)
    doc.add_picture(chart_margins(), width=Inches(6))
    doc.paragraphs[-1].add_run()  # spacer

    # ── Pages 4-5: Segments ──
    doc.add_page_break()
    add_heading(doc, "Revenue by Segment", level=2)
    doc.add_picture(chart_segments(), width=Inches(5))

    seg_tbl = doc.add_table(rows=1, cols=3)
    seg_tbl.style = 'Table Grid'
    for i, h in enumerate(['Segment', 'Revenue', 'YoY']):
        seg_tbl.rows[0].cells[i].text = h
        seg_tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for s in SEGMENTS:
        r = seg_tbl.add_row().cells
        r[0].text = s[0]
        r[1].text = f"{CURRENCY} {s[1]:.1f}B"
        r[2].text = f"{'+' if s[2] >= 0 else ''}{s[2]}%"
    doc.add_paragraph()

    # ── Pages 6-7: Investment Thesis ──
    doc.add_page_break()
    add_heading(doc, "Investment Thesis Update", level=2)
    for t in THESIS:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{t[0]} {t[1]}: ").bold = True
        p.add_run(t[2])

    # ── Pages 8-10: Estimates & Valuation ──
    doc.add_page_break()
    add_heading(doc, "Estimate Revisions (FY2026E)", level=2)
    doc.add_picture(chart_estimate_revisions(), width=Inches(6))

    est_tbl = doc.add_table(rows=1, cols=4)
    est_tbl.style = 'Table Grid'
    for i, h in enumerate(['Metric', 'Old Est', 'New Est', 'Change']):
        est_tbl.rows[0].cells[i].text = h
        est_tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for e in EST_REVISIONS:
        r = est_tbl.add_row().cells
        chg = (e[2] - e[1]) / e[1] * 100
        r[0].text = f"{e[0]} ({e[3]})"
        r[1].text = str(e[1])
        r[2].text = str(e[2])
        r[3].text = f"{chg:+.1f}%"
    doc.add_paragraph()

    # ── Risks ──
    add_heading(doc, "Key Risks", level=2)
    doc.add_paragraph("Risks: " + " · ".join(f"`{r}`" for r in RISKS))

    doc.save(OUTPUT_FILE)
    print(f"Saved: {OUTPUT_FILE}")


if __name__ == '__main__':
    build_docx()