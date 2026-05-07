"""
earnings-preview DOCX builder
==============================
Reusable helper library for generating earnings preview DOCX reports.

Usage
-----
    from scripts.docx_builder import DocxBuilder

    b = DocxBuilder(
        symbol="ADBE.US",
        company="Adobe Inc.",
        report_date="2026年6月11日",
        analysis_date="2026年5月6日",
        price="$250.71",
        market_cap="~$1013亿",
        valuation="P/E 14.06x",
        rating="买入（39位分析师）",
        output_path="ADBE_Q2_FY2026_Earnings_Preview.docx",
    )

    b.cover()
    b.toc(sections)       # list of (number, title) tuples
    b.section("【一】上期业绩指引回顾")
    b.body("说明文字")
    b.bullet("要点")
    b.table(headers, rows, col_widths)
    b.qa("问题", "答复", "本次关注")
    b.section("【八】风险提示")
    ...
    b.disclaimer()
    b.save()

Font policy
-----------
- Latin / numbers  : Calibri
- CJK characters   : Microsoft YaHei
- Every run sets BOTH fonts so tables and headings render correctly on all platforms.
- To override, pass latin= or cjk= to any method.
"""

from __future__ import annotations

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── Colour palette ────────────────────────────────────────────────────────────
BLUE = RGBColor(0x1A, 0x56, 0xDB)
DARK = RGBColor(0x37, 0x40, 0x51)
GREY = RGBColor(0x6B, 0x72, 0x80)
HEAD_FILL = "DBE9FF"

DEFAULT_LATIN = "Calibri"
DEFAULT_CJK   = "Microsoft YaHei"


# ── Low-level font helper ─────────────────────────────────────────────────────
def _set_run_fonts(run, latin: str, cjk: str) -> None:
    """Set both Latin and CJK fonts on a run element."""
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rPr.get_or_add_rFonts().set(qn("w:eastAsia"), cjk)


def _set_doc_default_cjk(doc: Document, cjk: str = DEFAULT_CJK) -> None:
    """Apply CJK font to the document's Normal style so unstyled runs inherit it."""
    for name in ("Normal", "Default Paragraph Font"):
        try:
            rPr = doc.styles[name].element.get_or_add_rPr()
            rPr.get_or_add_rFonts().set(qn("w:eastAsia"), cjk)
        except Exception:
            pass


# ── Builder class ─────────────────────────────────────────────────────────────
# ── Language label tables ──────────────────────────────────────────────────────
_LABELS = {
    "zh": {
        "subtitle":      "财报前瞻",
        "analysis_date": "分析日期",
        "report_date":   "财报发布日期",
        "price":         "现价",
        "market_cap":    "市值",
        "rating":        "评级",
        "toc":           "目录",
        "qa_answer":     "  \u21b3 管理层答复：",
        "qa_watch":      "  \u21b3 本次关注：",
        "disclaimer_h":  "免责声明",
        "disclaimer_b":  (
            "本报告仅供参考和教育用途，不构成任何投资建议、买卖证券的邀约或任何形式的推荐。"
            "数据来源包括 Longbridge CLI、SEC 公开文件及互联网公开信息。"
            "所有估算、情景分析及预测均存在重大不确定性，过往表现不代表未来结果。"
            "投资者在作出任何投资决策前应进行独立尽职调查。"
        ),
    },
    "en": {
        "subtitle":      "Earnings Preview",
        "analysis_date": "Analysis Date",
        "report_date":   "Report Date",
        "price":         "Price",
        "market_cap":    "Mkt Cap",
        "rating":        "Rating",
        "toc":           "Table of Contents",
        "qa_answer":     "  \u21b3 Management Response: ",
        "qa_watch":      "  \u21b3 Watch For: ",
        "disclaimer_h":  "Disclaimer",
        "disclaimer_b":  (
            "This report is for informational and educational purposes only and does not "
            "constitute investment advice, a solicitation to buy or sell securities, or any "
            "form of recommendation. Data sources include the Longbridge CLI, SEC public "
            "filings, and publicly available internet information. All estimates, scenario "
            "analyses, and forecasts involve material uncertainty; past performance does not "
            "guarantee future results. Investors should conduct independent due diligence "
            "before making any investment decision."
        ),
    },
}


class DocxBuilder:
    def __init__(
        self,
        symbol: str,
        company: str,
        report_date: str,
        analysis_date: str,
        price: str,
        market_cap: str,
        valuation: str,
        rating: str,
        output_path: str,
        lang: str = "zh",
        latin: str = DEFAULT_LATIN,
        cjk: str = DEFAULT_CJK,
    ) -> None:
        self.symbol        = symbol
        self.company       = company
        self.report_date   = report_date
        self.analysis_date = analysis_date
        self.price         = price
        self.market_cap    = market_cap
        self.valuation     = valuation
        self.rating        = rating
        self.output_path   = output_path
        self.lang          = lang if lang in _LABELS else "zh"
        self.latin         = latin
        self.cjk           = cjk
        self._lbl          = _LABELS[self.lang]

        self.doc = Document()
        sec = self.doc.sections[0]
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.2)
        sec.right_margin  = Inches(1.2)

        _set_doc_default_cjk(self.doc, cjk)

    # ── Internal run factory ──────────────────────────────────────────────────
    def _run(self, para, text: str, bold: bool = False, size: int = 10,
             color: RGBColor | None = None,
             latin: str | None = None, cjk: str | None = None):
        run = para.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(size)
        _set_run_fonts(run, latin or self.latin, cjk or self.cjk)
        if color:
            run.font.color.rgb = color
        return run

    # ── Page elements ─────────────────────────────────────────────────────────
    def page_break(self):
        self.doc.add_page_break()

    def hr(self):
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "C0C0C0")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── Cover page ────────────────────────────────────────────────────────────
    def cover(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(48)
        self._run(p, f"{self.company}（{self.symbol}）",
                  bold=True, size=22, color=BLUE)

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._run(p2, self._lbl["subtitle"], bold=True, size=16)

        self.doc.add_paragraph()

        p3 = self.doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._run(p3,
                  f"{self._lbl['analysis_date']}：{self.analysis_date}    |    "
                  f"{self._lbl['report_date']}：{self.report_date}",
                  size=11, color=GREY)

        self.doc.add_paragraph()

        p4 = self.doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._run(p4,
                  f"{self._lbl['price']}：{self.price}  |  {self._lbl['market_cap']}：{self.market_cap}  |  "
                  f"{self.valuation}  |  {self._lbl['rating']}：{self.rating}",
                  size=10)

        self.doc.add_page_break()

    # ── Table of contents ─────────────────────────────────────────────────────
    def toc(self, sections: list[tuple[str, str]]):
        """sections: list of (number_label, title), e.g. [('【一】', '上期业绩指引回顾'), ...]"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after  = Pt(4)
        self._run(p, self._lbl["toc"], bold=True, size=15, color=BLUE)

        for num, title in sections:
            row = self.doc.add_paragraph()
            row.paragraph_format.space_before = Pt(3)
            row.paragraph_format.space_after  = Pt(3)
            self._run(row, num + "  ", bold=True, size=10, color=BLUE)
            self._run(row, title, size=10)

        self.doc.add_page_break()

    # ── Section heading ───────────────────────────────────────────────────────
    def section(self, text: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after  = Pt(4)
        self._run(p, text, bold=True, size=15, color=BLUE)

    def subsection(self, text: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        self._run(p, text, bold=True, size=12, color=DARK)

    # ── Body text ─────────────────────────────────────────────────────────────
    def body(self, text: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        self._run(p, text, size=10)

    def bullet(self, text: str):
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        self._run(p, text, size=10)

    # ── Q&A block ─────────────────────────────────────────────────────────────
    def qa(self, question: str, answer: str, watch: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        self._run(p, question, bold=True, size=10)

        for label, val in [(self._lbl["qa_answer"], answer), (self._lbl["qa_watch"], watch)]:
            p2 = self.doc.add_paragraph()
            p2.paragraph_format.left_indent  = Inches(0.3)
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after  = Pt(3)
            self._run(p2, label, bold=True, size=9, color=DARK)
            self._run(p2, val, size=9)

    # ── Table ─────────────────────────────────────────────────────────────────
    def table(self, headers: list[str], rows: list[list],
              col_widths: list[float] | None = None):
        t = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(h))
            run.bold = True
            run.font.size = Pt(9)
            _set_run_fonts(run, self.latin, self.cjk)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  HEAD_FILL)
            cell._tc.get_or_add_tcPr().append(shd)

        # Data rows
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                cell = t.rows[r_i + 1].cells[c_i]
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run(str(val))
                run.font.size = Pt(9)
                _set_run_fonts(run, self.latin, self.cjk)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        if col_widths:
            for i, w in enumerate(col_widths):
                for row in t.rows:
                    row.cells[i].width = Inches(w)

        self.doc.add_paragraph()

    # ── Image ─────────────────────────────────────────────────────────────────
    def image(self, path: str, width: float = 6.0, caption: str = ""):
        """Embed a PNG/JPG chart into the document. width is in inches."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        if caption:
            cp = self.doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(0)
            cp.paragraph_format.space_after  = Pt(6)
            self._run(cp, caption, size=8, color=GREY)
        else:
            self.doc.add_paragraph()

    # ── Disclaimer ────────────────────────────────────────────────────────────
    def disclaimer(self):
        self.hr()
        self.subsection(self._lbl["disclaimer_h"])
        self.body(self._lbl["disclaimer_b"])

    # ── Save ─────────────────────────────────────────────────────────────────
    def save(self) -> str:
        self.doc.save(self.output_path)
        return self.output_path
